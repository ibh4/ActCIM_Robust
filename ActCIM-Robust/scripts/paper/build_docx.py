"""Convert reports/final/ActCIM_Robust_paper_final.md into a typeset DOCX.

- Headings -> Word heading styles (navigable outline)
- Markdown tables -> Word tables (Table Grid, bold header)
- Display math $$...$$ -> 300-DPI matplotlib-mathtext PNG, centered
- Inline math $...$   -> Unicode text (greek letters / sub / superscripts)
- Images ![..](..)    -> embedded 300-DPI PNGs (paths resolved vs. the MD file)
- Bold / inline code  -> proper runs; CJK font Songti SC, latin Times New Roman

The DOCX is then converted to PDF with LibreOffice (done by the caller).
"""
from __future__ import annotations

import re
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from matplotlib import font_manager
from PIL import Image

ROOT = Path(__file__).resolve().parents[2]
MD = ROOT / "reports" / "final" / "ActCIM_Robust_paper_final.md"
OUT_DOCX = ROOT / "reports" / "final" / "ActCIM_Robust_paper_final.docx"
EQ_DIR = ROOT / "reports" / "final" / "_eq_png"
EQ_DIR.mkdir(exist_ok=True)

for fp in ["/System/Library/Fonts/Songti.ttc", "/System/Library/Fonts/STHeiti Medium.ttc"]:
    if Path(fp).exists():
        try:
            font_manager.fontManager.addfont(fp)
        except Exception:
            pass

EA_FONT = "Songti SC"
LATIN_FONT = "Times New Roman"

# ------------------------------------------------------------- latex -> text
GREEK = {
    r"\alpha": "α", r"\beta": "β", r"\lambda": "λ", r"\sigma": "σ",
    r"\Sigma": "Σ", r"\pi": "π", r"\mu": "μ", r"\Delta": "Δ",
}
SUP = str.maketrans("0123456789+-=()n", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿ")
SUB = str.maketrans("0123456789+-=()aeioruvxlmng", "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₐₑᵢₒᵣᵤᵥₓₗₘₙ₉")


def _try_script(body: str, table) -> str | None:
    out = body.translate(table)
    # only accept if every char actually converted
    if all(ord(c) > 0x2000 or not c.isascii() for c in out):
        return out
    return None


def latex_to_text(s: str) -> str:
    """Best-effort inline LaTeX -> Unicode plain text."""
    s = s.strip()
    for k, v in GREEK.items():
        s = s.replace(k, v)
    s = s.replace(r"\tilde{y}", "ỹ").replace(r"\tilde y", "ỹ")
    s = re.sub(r"\\(?:mathrm|text|mathbf|operatorname)\{([^{}]*)\}", r"\1", s)
    s = s.replace(r"\left", "").replace(r"\right", "")
    s = s.replace(r"\arg\min", "argmin").replace(r"\argmin", "argmin")
    s = s.replace(r"\max", "max").replace(r"\min", "min")
    s = s.replace(r"\sum", "Σ").replace(r"\int", "∫")
    s = s.replace(r"\in", "∈").replace(r"\sim", "~")
    s = s.replace(r"\neq", "≠").replace(r"\leq", "≤").replace(r"\geq", "≥")
    s = s.replace(r"\approx", "≈").replace(r"\pm", "±")
    s = s.replace(r"\cdot", "·").replace(r"\times", "×")
    s = s.replace(r"\qquad", "  ").replace(r"\quad", " ")
    s = s.replace(r"\,", "").replace(r"\;", " ").replace(r"\!", "")
    s = s.replace(r"\ ", " ").replace(r"\|", "‖")
    # \frac{a}{b} -> a/b (repeat for nesting)
    for _ in range(4):
        s2 = re.sub(r"\\frac\{([^{}]*)\}\{([^{}]*)\}", r"\1/\2", s)
        if s2 == s:
            break
        s = s2
    # super / subscripts
    def sup_repl(m):
        body = m.group(1) or m.group(2)
        conv = _try_script(body, SUP)
        return conv if conv else "^(" + body + ")"

    def sub_repl(m):
        body = m.group(1) or m.group(2)
        if body in ("max", "min"):
            return "_" + body
        conv = _try_script(body, SUB)
        return conv if conv else "_" + body

    s = re.sub(r"\^\{([^{}]*)\}|\^(\w)", sup_repl, s)
    s = re.sub(r"_\{([^{}]*)\}|_(\w)", sub_repl, s)
    s = s.replace("{", "").replace("}", "").replace("\\", "")
    return s


# -------------------------------------------------- display math -> PNG image
_EQ_COUNT = 0


def sanitize_for_mathtext(latex: str) -> tuple[str, str]:
    """Return (mathtext-safe latex, trailing plain-text note)."""
    note = ""
    # pull out any \text{...} groups that contain CJK -> plain trailing note
    idx = latex.find(r"\ \text{")
    if idx == -1:
        m = re.search(r"\\text\{[^}]*[\u4e00-\u9fff]", latex)
        idx = m.start() if m else -1
    if idx != -1:
        note = latex_to_text(latex[idx:])
        latex = latex[:idx].rstrip().rstrip(",").rstrip("\\")
    latex = latex.replace(r"\qquad", r"\quad\quad")
    latex = latex.replace(r"\text{", r"\mathrm{")
    latex = latex.replace(r"\max", r"\mathrm{max}")
    latex = latex.replace(r"\min", r"\mathrm{min}")
    return latex, note


def render_equation(latex: str) -> tuple[Path, float, str]:
    """Render display latex via mathtext; return (png, width_in_inches, note)."""
    global _EQ_COUNT
    _EQ_COUNT += 1
    safe, note = sanitize_for_mathtext(latex)
    out = EQ_DIR / f"eq{_EQ_COUNT:02d}.png"
    fig = plt.figure(figsize=(0.1, 0.1))
    t = fig.text(0, 0, f"${safe}$", fontsize=13)
    fig.canvas.draw()
    bbox = t.get_window_extent()
    fig.set_size_inches(bbox.width / fig.dpi + 0.05, bbox.height / fig.dpi + 0.05)
    fig.savefig(out, dpi=300, transparent=True, bbox_inches="tight", pad_inches=0.02)
    plt.close(fig)
    with Image.open(out) as im:
        w_px = im.width
    return out, w_px / 300.0, note


# ------------------------------------------------------------- docx helpers
def set_run(run, size=10.5, bold=False, italic=False, code=False,
            color=None, sub=False, sup=False):
    run.font.name = "Courier New" if code else LATIN_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), EA_FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    if sub:
        run.font.subscript = True
    if sup:
        run.font.superscript = True


TOKEN_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\$[^$]+\$)")


def add_rich_text(par, text: str, size=10.5, base_bold=False, color=None):
    for tok in TOKEN_RE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            inner = tok[2:-2]
            # bold section may itself contain $math$
            for t2 in re.split(r"(\$[^$]+\$)", inner):
                if not t2:
                    continue
                if t2.startswith("$") and t2.endswith("$"):
                    r = par.add_run(latex_to_text(t2[1:-1]))
                    set_run(r, size=size, bold=True, italic=True, color=color)
                else:
                    r = par.add_run(t2)
                    set_run(r, size=size, bold=True, color=color)
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1])
            set_run(r, size=size - 0.5, code=True, color=(90, 60, 60))
        elif tok.startswith("$") and tok.endswith("$"):
            r = par.add_run(latex_to_text(tok[1:-1]))
            set_run(r, size=size, italic=True, bold=base_bold, color=color)
        else:
            r = par.add_run(tok)
            set_run(r, size=size, bold=base_bold, color=color)


def add_heading(doc, text: str, level: int):
    par = doc.add_heading("", level=level)
    sizes = {0: 16, 1: 14, 2: 12, 3: 11}
    add_rich_text(par, text, size=sizes.get(level, 11), base_bold=True,
                  color=(0, 0, 0))
    if level == 0:
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return par


def add_table(doc, header: list[str], rows: list[list[str]]):
    tab = doc.add_table(rows=len(rows) + 1, cols=len(header))
    tab.style = "Table Grid"
    tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, cell_text in enumerate(header):
        cell = tab.rows[0].cells[j]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_rich_text(cell.paragraphs[0], cell_text, size=9, base_bold=True)
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= len(header):
                continue
            cell = tab.rows[i + 1].cells[j]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_rich_text(cell.paragraphs[0], cell_text, size=9)
    doc.add_paragraph()


# ------------------------------------------------------------------- build
def build():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.4)
        sec.bottom_margin = Cm(2.4)
        sec.left_margin = Cm(2.6)
        sec.right_margin = Cm(2.6)
    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), EA_FONT)

    lines = MD.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        # ---- headings
        if stripped.startswith("#"):
            hashes = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[hashes:].strip()
            add_heading(doc, text, level=min(hashes - 1, 3) if hashes > 1 else 0)
            i += 1
            continue
        # ---- blockquote
        if stripped.startswith(">"):
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Cm(0.8)
            add_rich_text(par, stripped.lstrip("> "), size=9, color=(96, 96, 96))
            i += 1
            continue
        # ---- display math
        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
            png, w_in, note = render_equation(stripped[2:-2])
            par = doc.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = par.add_run()
            run.add_picture(str(png), width=Inches(min(w_in, 6.0)))
            if note:
                r = par.add_run("  " + note)
                set_run(r, size=9.5)
            i += 1
            continue
        # ---- image
        m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", stripped)
        if m:
            img = (MD.parent / m.group(2)).resolve()
            par = doc.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if img.exists():
                par.add_run().add_picture(str(img), width=Cm(15.2))
            else:
                add_rich_text(par, f"[缺失图片: {m.group(2)}]", size=9)
            i += 1
            continue
        # ---- table
        if stripped.startswith("|") and i + 1 < len(lines) and \
                re.match(r"^\|[\s:\-|]+\|$", lines[i + 1].strip()):
            header = [c.strip() for c in stripped.strip("|").split("|")]
            rows = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            add_table(doc, header, rows)
            continue
        # ---- figure caption (**图N** ...)
        if re.match(r"^\*\*图\d+\*\*", stripped):
            par = doc.add_paragraph()
            par.paragraph_format.space_after = Pt(10)
            add_rich_text(par, stripped, size=9, color=(70, 70, 70))
            i += 1
            continue
        # ---- bullet list
        if stripped.startswith("- "):
            par = doc.add_paragraph(style="List Bullet")
            add_rich_text(par, stripped[2:])
            i += 1
            continue
        # ---- numbered list
        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m and not stripped.startswith("10.1"):
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Cm(0.6)
            add_rich_text(par, f"{m.group(1)}. {m.group(2)}")
            i += 1
            continue
        # ---- reference entries [n] ...
        if re.match(r"^\[\d+\]", stripped):
            par = doc.add_paragraph()
            par.paragraph_format.space_after = Pt(4)
            add_rich_text(par, stripped, size=9.5)
            i += 1
            continue
        # ---- normal paragraph
        par = doc.add_paragraph()
        par.paragraph_format.first_line_indent = Cm(0.74)
        par.paragraph_format.space_after = Pt(4)
        par.paragraph_format.line_spacing = 1.35
        add_rich_text(par, stripped)
        i += 1

    doc.save(OUT_DOCX)
    print("saved", OUT_DOCX)


if __name__ == "__main__":
    build()
